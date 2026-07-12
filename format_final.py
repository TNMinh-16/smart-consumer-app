from pathlib import Path
from copy import deepcopy
from zipfile import ZipFile
from lxml import etree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "final.docx"
OUT = ROOT / "final_da_chinh_sua.docx"


def set_run_font(run, name="Times New Roman", size=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, name="Times New Roman", size=None, bold=None):
    style.font.name = name
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_style_spacing(style, line=1.5, before=0, after=6, first_cm=1.0, left_cm=0):
    pf = style.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(first_cm) if first_cm else None
    pf.left_indent = Cm(left_cm) if left_cm else None


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_update_fields(doc):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def style_toc_parts(doc):
    # These styles already exist in the source file, but python-docx does not
    # expose them by their UI names. Configure them through the style table.
    styles = doc.styles
    for name, size, bold, left, first in (
        ("TOCHeading", 14, True, 0, 0),
        ("TOC1", 13, False, 0, 0),
        ("TOC2", 13, False, 0.65, 0),
        ("TOC3", 13, False, 1.3, 0),
    ):
        try:
            st = styles[name]
        except KeyError:
            continue
        set_style_font(st, size=size, bold=bold)
        pf = st.paragraph_format
        pf.line_spacing = 1.15
        pf.space_before = Pt(0)
        pf.space_after = Pt(2)
        pf.left_indent = Cm(left) if left else None
        pf.first_line_indent = Cm(first) if first else None


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=13)
    set_style_spacing(normal, line=1.5, before=0, after=6, first_cm=1.0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    lp = doc.styles["List Paragraph"]
    set_style_font(lp, size=13)
    lp.paragraph_format.line_spacing = 1.5
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(3)
    lp.paragraph_format.left_indent = Cm(0.75)
    lp.paragraph_format.first_line_indent = Cm(-0.5)
    lp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after in (
        ("Heading 1", 16, 12, 6),
        ("Heading 2", 14, 10, 4),
        ("Heading 3", 13, 8, 3),
    ):
        st = doc.styles[name]
        set_style_font(st, size=size, bold=True)
        set_style_spacing(st, line=1.5, before=before, after=after, first_cm=0)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.keep_together = True

    style_toc_parts(doc)


def heading_level(text):
    t = text.strip()
    if t in {"MỞ ĐẦU", "NỘI DUNG CHÍNH", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO"}:
        return 1
    if t.startswith(("I.1.", "I.2.", "I.3.", "II.1.", "II.2.", "II.3.")):
        # The first numbered component after the roman numeral is level 2.
        parts = t.split(".", 3)
        if len(parts) >= 3 and parts[0] == "I" and parts[1] == "2" and parts[2].startswith(("1", "2")):
            return 3
        if len(parts) >= 3 and parts[0] == "I" and parts[1] == "3" and parts[2].startswith(("1", "2", "3")):
            return 3
        if len(parts) >= 3 and parts[0] == "II" and parts[1] == "3" and parts[2].startswith(("1", "2", "3")):
            return 3
        return 2
    return 0


def snapshot_content(doc):
    paras = [p.text for p in doc.paragraphs]
    tables = [[[c.text for c in row.cells] for row in table.rows] for table in doc.tables]
    return paras, tables


def apply_layout(doc):
    # A4, Vietnamese academic-paper margins.
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # The first 17 direct paragraphs are the cover; preserve their hierarchy
    # while keeping their existing centered/bold emphasis.
    for idx, p in enumerate(doc.paragraphs):
        if idx < 17:
            for r in p.runs:
                set_run_font(r, size=(r.font.size.pt if r.font.size else None))
            continue

        level = heading_level(p.text)
        if level:
            p.style = doc.styles[f"Heading {level}"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = None
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            size = {1: 16, 2: 14, 3: 13}[level]
            for r in p.runs:
                set_run_font(r, size=size)
                r.bold = True
        else:
            if p.style.name not in {"List Paragraph"}:
                p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if p.text.strip() == "Link":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = None
            for r in p.runs:
                set_run_font(r, size=13)

    # Table text stays unchanged but receives a readable academic-paper table style.
    for ti, table in enumerate(doc.tables):
        table.autofit = False
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = None
                    for r in p.runs:
                        set_run_font(r, size=10.5)
                        if ri == 0:
                            r.bold = True

    # Keep the existing section/page breaks; fields will be refreshed in Word.
    add_update_fields(doc)


def main():
    doc = Document(str(SRC))
    before = snapshot_content(doc)
    configure_styles(doc)
    apply_layout(doc)
    doc.save(str(OUT))

    check = Document(str(OUT))
    after = snapshot_content(check)
    if before != after:
        raise RuntimeError("Nội dung ngoài mục lục đã thay đổi ngoài ý muốn")
    print(f"Wrote {OUT}")
    print("Content check: unchanged")


if __name__ == "__main__":
    main()
